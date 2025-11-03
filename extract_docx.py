from docx import Document
import sys
import json

def extract_docx(file_path):
    doc = Document(file_path)
    content = {
        'paragraphs': [],
        'tables': []
    }

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            content['paragraphs'].append({
                'text': para.text,
                'style': para.style.name
            })

    # Extract tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data.append(row_data)
        content['tables'].append(table_data)

    return content

if __name__ == '__main__':
    file_path = sys.argv[1]
    content = extract_docx(file_path)
    print(json.dumps(content, ensure_ascii=False, indent=2))
